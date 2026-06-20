"""
Warrior Prep — Student Workbook Word Document Builder
Run: python3 warrior_prep_docx.py
Generates: output/warrior_prep_student_workbook.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from content import UNITS, SCHOOL_NAME, COURSE_TITLE, GRADE

os.makedirs("output", exist_ok=True)

WARRIOR_RED  = RGBColor(0xC7, 0x14, 0x14)
WARRIOR_DARK = RGBColor(0x21, 0x21, 0x21)
WARRIOR_GOLD = RGBColor(0xF5, 0xE0, 0xB2)
LIGHT_BLUE   = RGBColor(0xE8, 0xF0, 0xFE)
CREAM_BG     = RGBColor(0xFD, 0xFA, 0xF4)
LIGHT_GRAY   = RGBColor(0xF5, 0xF5, 0xF5)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(shd)


def set_cell_borders(cell, top=False, bottom=False, left=False, right=False,
                     color="AAAAAA", size="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, on in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if on:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_paragraph(doc_or_cell, text, bold=False, italic=False, size=11,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs else doc_or_cell.add_paragraph()
    p.alignment = align
    pPr = p._p.get_or_add_pPr()
    spA = OxmlElement("w:spacing")
    spA.set(qn("w:after"), str(int(space_after * 20)))
    pPr.append(spA)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_heading_bar(doc, text, color: RGBColor = None, size=12):
    color = color or WARRIOR_RED
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = WHITE
    cell._tc.get_or_add_tcPr()
    doc.add_paragraph("")


def add_ruled_lines(doc, n=4, width_inches=6.0):
    for _ in range(n):
        t = doc.add_table(rows=1, cols=1)
        cell = t.rows[0].cells[0]
        cell.width = Inches(width_inches)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        el = OxmlElement("w:bottom")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "C8D0DC")
        tcBorders.append(el)
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        pPr = p._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:line"), "360")
        sp.set(qn("w:lineRule"), "exact")
        pPr.append(sp)


def color_from_tuple(t):
    if isinstance(t, (list, tuple)) and len(t) == 3:
        return RGBColor(int(t[0]*255), int(t[1]*255), int(t[2]*255))
    return WARRIOR_RED


# ── Sections ──────────────────────────────────────────────────────────────────

def add_cover(doc):
    doc.add_paragraph("")
    doc.add_paragraph("")
    h = doc.add_heading(COURSE_TITLE, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = WARRIOR_RED
        run.font.size = Pt(36)

    p = doc.add_paragraph("Student Workbook")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(18)
    p.runs[0].font.color.rgb = WARRIOR_DARK

    p2 = doc.add_paragraph(f"{GRADE}  ·  Semester Elective")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].font.size = Pt(14)

    p3 = doc.add_paragraph(SCHOOL_NAME)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].font.size = Pt(12)
    p3.runs[0].font.color.rgb = WARRIOR_DARK

    doc.add_paragraph("")
    doc.add_paragraph("")

    for label in ["Name: _______________________________   Period: _______",
                  "Teacher: _____________________________   Year: _________"]:
        p = doc.add_paragraph(label)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()


def add_how_to_use(doc):
    add_heading_bar(doc, "HOW TO USE THIS WORKBOOK")
    h = doc.add_heading("Welcome to Warrior Prep!", level=2)
    for run in h.runs:
        run.font.color.rgb = WARRIOR_RED

    doc.add_paragraph(
        "This workbook is your guide, notebook, and planner for the entire semester. "
        "Here is what's inside and how to get the most out of it."
    )

    items = [
        ("Planner Section (Weeks 1–18)",
         "One page per week. Use it for ALL six of your courses — every day."),
        ("Cornell Notes Pages",
         "Every lesson has a Cornell Notes page: left column = questions, right = notes, bottom = summary."),
        ("Vocabulary & Morphology",
         "Each lesson has a Word Focus box with etymology (word roots) to help you remember meaning."),
        ("Exit Tickets",
         "Complete the exit ticket at the end of each lesson. Formats vary throughout the course."),
        ("Pre-Tests & Post-Tests",
         "Each unit starts with a pre-test (baseline only) and ends with a post-test to show your growth."),
    ]
    for title, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title}:  ")
        run.bold = True
        run.font.color.rgb = WARRIOR_RED
        p.add_run(desc)

    doc.add_paragraph("")
    add_heading_bar(doc, "AVID CORNELL NOTES — HOW THEY WORK", color=WARRIOR_DARK)

    t = doc.add_table(rows=2, cols=3)
    t.style = "Table Grid"
    hdrs = ["LEFT COLUMN (Questions/Cues)", "RIGHT COLUMN (Notes)", "BOTTOM (Summary)"]
    descs = [
        "Essential questions, key terms, cues.\n\nFill in DURING or right AFTER notes.",
        "Main notes, facts, ideas, examples.\nIn your own words.\nLeave space to add later.",
        "3–5 sentences after the lesson.\n\nNO peeking! Your words only.",
    ]
    for i, (hdr, desc) in enumerate(zip(hdrs, descs)):
        cell = t.rows[0].cells[i]
        set_cell_bg(cell, WARRIOR_RED)
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        cell2 = t.rows[1].cells[i]
        set_cell_bg(cell2, LIGHT_GRAY)
        cell2.paragraphs[0].add_run(desc).font.size = Pt(9)

    doc.add_page_break()


def add_planner(doc):
    add_heading_bar(doc, "WARRIOR PREP — SEMESTER PLANNER  (Weeks 1–18)")
    h = doc.add_heading("Your Semester Planner", level=2)
    for run in h.runs:
        run.font.color.rgb = WARRIOR_RED
    doc.add_paragraph(
        "Use ONE page per week. On Week 1, write your six course names in the Course column. "
        "Record every assignment, test, and deadline on the day it is ASSIGNED."
    )
    for tip in ["Write the assignment and due date.", "✓ = done   → = carried forward",
                "Color-code by subject if helpful.", "Keep this with you all day."]:
        p = doc.add_paragraph(f"•  {tip}", style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
    for wk in range(1, 19):
        h = doc.add_heading(f"Week {wk}  ·  Dates: _______________________", level=3)
        h.runs[0].font.color.rgb = WARRIOR_DARK

        t = doc.add_table(rows=7, cols=6)
        t.style = "Table Grid"
        # Header row
        set_cell_bg(t.rows[0].cells[0], WARRIOR_DARK)
        t.rows[0].cells[0].paragraphs[0].add_run("Course").font.color.rgb = WHITE

        for j, day in enumerate(days, 1):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(day).font.color.rgb = WHITE
            t.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 6 course rows
        for i in range(1, 7):
            label = f"Course {i}: ___________" if wk == 1 else f"Course {i}"
            t.rows[i].cells[0].paragraphs[0].add_run(label).font.size = Pt(8)
            set_cell_bg(t.rows[i].cells[0], LIGHT_GRAY)
            for j in range(1, 6):
                bg_clr = CREAM_BG if i % 2 == 1 else WHITE
                set_cell_bg(t.rows[i].cells[j], bg_clr)
                # Set row height via XML
                tr = t.rows[i]._tr
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement("w:trHeight")
                trHeight.set(qn("w:val"), "700")  # ~0.5 inch in twips
                trPr.append(trHeight)

        # Goals row
        doc.add_paragraph("")
        gt = doc.add_table(rows=2, cols=1)
        gt.style = "Table Grid"
        set_cell_bg(gt.rows[0].cells[0], WARRIOR_GOLD)
        r = gt.rows[0].cells[0].paragraphs[0].add_run("Weekly Goals / Notes / Carry-Forward")
        r.bold = True
        r.font.size = Pt(9)
        set_cell_bg(gt.rows[1].cells[0], WHITE)
        for _ in range(3):
            gt.rows[1].cells[0].add_paragraph("")

        doc.add_page_break()


def add_unit_title(doc, unit):
    uc = color_from_tuple(unit["color"])
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, uc)
    p = cell.paragraphs[0]
    r = p.add_run(f"Unit {unit['number']}  ·  {unit['name']}")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = WHITE
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = cell.add_paragraph(unit["overview"])
    p2.runs[0].font.size = Pt(12)
    p2.runs[0].font.color.rgb = WHITE

    cell.add_paragraph("")
    ph = cell.add_paragraph("Unit Vocabulary")
    ph.runs[0].bold = True
    ph.runs[0].font.color.rgb = WARRIOR_GOLD
    ph.runs[0].font.size = Pt(11)

    for v in unit["unit_vocabulary"]:
        pv = cell.add_paragraph()
        rv = pv.add_run(f"{v['word']}: ")
        rv.bold = True
        rv.font.color.rgb = WARRIOR_GOLD
        rv.font.size = Pt(10)
        pv.add_run(f"{v['definition']}  [Etymology: {v['etymology']}]").font.size = Pt(9)

    doc.add_page_break()


def add_test_page(doc, unit, is_pretest=True):
    uc = color_from_tuple(unit["color"])
    label = "PRE-TEST" if is_pretest else "POST-TEST"
    add_heading_bar(doc, f"Unit {unit['number']}: {unit['name']}  —  {label}", color=uc)
    if is_pretest:
        doc.add_paragraph(
            "Directions: Answer each question as completely as you can right now. "
            "This is NOT graded for correctness — it shows your starting point. Be honest!"
        )
        for i, q in enumerate(unit["pretest"], 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}.  ").bold = True
            p.add_run(q)
            add_ruled_lines(doc, 4)
    else:
        doc.add_paragraph(f"Directions: {unit['posttest_prompt']}")
        add_ruled_lines(doc, 20)
    doc.add_page_break()


def add_cornell_page(doc, unit, lesson):
    uc = color_from_tuple(unit["color"])
    sc = lesson["scaffolding"]
    vf = lesson["vocab_focus"]

    add_heading_bar(doc, f"Lesson {lesson['lesson_label']}  —  {lesson['title']}  ·  Cornell Notes", color=uc, size=11)

    # Meta
    p = doc.add_paragraph("Name: _________________________________   Date: ________________   Period: ______")
    p.runs[0].font.size = Pt(9)

    # Vocab box
    vt = doc.add_table(rows=1, cols=3)
    vt.style = "Table Grid"
    set_cell_bg(vt.rows[0].cells[0], LIGHT_BLUE)
    r0 = vt.rows[0].cells[0].paragraphs[0].add_run(f"Word Focus: {vf['word']}")
    r0.bold = True
    r0.font.size = Pt(9)
    set_cell_bg(vt.rows[0].cells[1], LIGHT_BLUE)
    vt.rows[0].cells[1].paragraphs[0].add_run(f"Etymology: {vf['etymology']}").font.size = Pt(8)
    set_cell_bg(vt.rows[0].cells[2], LIGHT_BLUE)
    vt.rows[0].cells[2].paragraphs[0].add_run(vf["definition"]).font.size = Pt(9)

    doc.add_paragraph("")

    # Cornell table (2 columns, questions | notes)
    ct = doc.add_table(rows=1, cols=2)
    ct.style = "Table Grid"
    qcell = ct.rows[0].cells[0]
    ncell = ct.rows[0].cells[1]

    # Set column widths
    for i, w in enumerate([1.8, 4.7]):
        ct.columns[i].width = Inches(w)

    set_cell_bg(qcell, LIGHT_GRAY)

    # Questions column
    qp = qcell.paragraphs[0]
    r = qp.add_run("QUESTIONS / CUES")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = uc

    if sc <= 2:
        for q in lesson["cornell_questions"]:
            p2 = qcell.add_paragraph(f"• {q}")
            p2.runs[0].font.size = Pt(8)
        qcell.add_paragraph("Your questions:")
        for _ in range(5):
            p3 = qcell.add_paragraph("")
            set_cell_borders(qcell, bottom=True)
    else:
        for q in lesson["cornell_questions"][:2]:
            p2 = qcell.add_paragraph(f"• {q}")
            p2.runs[0].font.size = Pt(8)
        for _ in range(8):
            qcell.add_paragraph("")

    # Notes column
    np_ = ncell.paragraphs[0]
    r2 = np_.add_run("NOTES")
    r2.bold = True
    r2.font.size = Pt(8)
    r2.font.color.rgb = WARRIOR_DARK

    prefilled = lesson.get("prefilled_notes", "")
    if sc == 1 and prefilled:
        pf = ncell.add_paragraph(prefilled)
        pf.runs[0].font.size = Pt(8)
        pf.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
        ncell.add_paragraph("Continue your notes below:")
        for _ in range(8):
            ncell.add_paragraph("")
    elif sc == 2 and prefilled:
        pf = ncell.add_paragraph(prefilled)
        pf.runs[0].font.size = Pt(8)
        pf.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
        for _ in range(6):
            ncell.add_paragraph("")
    else:
        for _ in range(18):
            ncell.add_paragraph("")

    # Summary box
    st = doc.add_table(rows=2, cols=1)
    st.style = "Table Grid"
    set_cell_bg(st.rows[0].cells[0], RGBColor(0xE8, 0xE0, 0xD0))
    sr = st.rows[0].cells[0].paragraphs[0].add_run(
        "SUMMARY — Write the main ideas in your own words (complete AFTER the lesson):")
    sr.bold = True
    sr.font.size = Pt(8)
    set_cell_bg(st.rows[1].cells[0], WHITE)
    for _ in range(4):
        st.rows[1].cells[0].add_paragraph("")

    doc.add_page_break()


def add_lesson_content(doc, unit, lesson):
    uc = color_from_tuple(unit["color"])
    add_heading_bar(doc, f"Unit {unit['number']}  ·  Lesson {lesson['lesson_label']}  —  {lesson['title']}", color=uc, size=11)

    # Objectives
    ot = doc.add_table(rows=1, cols=1)
    ot.style = "Table Grid"
    set_cell_bg(ot.rows[0].cells[0], LIGHT_BLUE)
    oc = ot.rows[0].cells[0]
    r = oc.paragraphs[0].add_run("Learning Objectives — By the end of this lesson you will be able to:")
    r.bold = True
    r.font.size = Pt(9)
    for obj in lesson["objectives"]:
        op = oc.add_paragraph(f"◉  {obj}")
        op.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    # Content sections
    for sec in lesson["content_sections"]:
        h = doc.add_heading(sec["heading"], level=3)
        h.runs[0].font.color.rgb = WARRIOR_DARK
        doc.add_paragraph(sec["body"])

        act = sec.get("activity")
        if act:
            _render_docx_activity(doc, act, uc)
        doc.add_paragraph("")


def _render_docx_activity(doc, act, uc):
    at = act["type"]

    if at in ("rating_table", "audit_table"):
        cols = act["columns"]
        t = doc.add_table(rows=1 + len(act["rows"]), cols=1 + len(cols))
        t.style = "Table Grid"
        set_cell_bg(t.rows[0].cells[0], WARRIOR_DARK)
        t.rows[0].cells[0].paragraphs[0].add_run("Statement" if at == "rating_table" else "Distractor").font.color.rgb = WHITE
        for j, c in enumerate(cols, 1):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(c).font.color.rgb = WHITE
        for i, row in enumerate(act["rows"], 1):
            t.rows[i].cells[0].paragraphs[0].add_run(row).font.size = Pt(9)
            for j in range(1, len(cols) + 1):
                t.rows[i].cells[j].paragraphs[0].add_run("☐").font.size = Pt(10)
        if at == "audit_table":
            doc.add_paragraph("My #1 distractor: _______________   My plan to address it:")
            add_ruled_lines(doc, 2)

    elif at == "matrix_sort":
        doc.add_paragraph("Draw the Eisenhower Matrix below (4 quadrants). Place each task in the correct quadrant.")
        add_ruled_lines(doc, 8)
        for i, task in enumerate(act["tasks"], 1):
            doc.add_paragraph(f"{i}. {task}").runs[0].font.size = Pt(9)

    elif at == "interest_table":
        t = doc.add_table(rows=1 + len(act["rows"]), cols=4)
        t.style = "Table Grid"
        for j, hdr in enumerate(["Career Field", "Description", "My Interest (1–5)", "Career I Know"]):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(hdr).font.color.rgb = WHITE
        for i, (name, desc) in enumerate(act["rows"], 1):
            t.rows[i].cells[0].paragraphs[0].add_run(name).font.size = Pt(9)
            t.rows[i].cells[1].paragraphs[0].add_run(desc).font.size = Pt(8)

    elif at == "two_column_chart":
        t = doc.add_table(rows=1 + len(act["fields"]), cols=3)
        t.style = "Table Grid"
        for j, hdr in enumerate(["", "Course 1", "Course 2"]):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(hdr).font.color.rgb = WHITE
        for i, field in enumerate(act["fields"], 1):
            t.rows[i].cells[0].paragraphs[0].add_run(field).font.size = Pt(9)

    elif at == "venn_diagram":
        doc.add_paragraph(f"Draw a Venn Diagram: {act['label_a']}  vs.  {act['label_b']}")
        add_ruled_lines(doc, 8)

    elif at == "guided_reflection":
        for p_text in act["prompts"]:
            doc.add_paragraph(p_text, style="List Bullet")
            add_ruled_lines(doc, 2)

    elif at == "smart_goal_template":
        for label, hint, lines in act["fields"]:
            p = doc.add_paragraph()
            p.add_run(f"{label}  ").bold = True
            p.add_run(f"({hint})").italic = True
            add_ruled_lines(doc, lines)

    elif at == "study_plan_template":
        for field in act["fields"]:
            doc.add_paragraph(field).runs[0].bold = True
            add_ruled_lines(doc, 2)

    elif at == "backward_plan":
        t = doc.add_table(rows=1 + act["weeks"], cols=3)
        t.style = "Table Grid"
        for j, hdr in enumerate(["Week", "Tasks to Complete", "My Mini-Deadline"]):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(hdr).font.color.rgb = WHITE
        for i, w in enumerate(range(act["weeks"], 0, -1), 1):
            t.rows[i].cells[0].paragraphs[0].add_run(f"Week {w}").font.size = Pt(9)

    elif at == "role_play":
        for sc in act["scenarios"]:
            p = doc.add_paragraph()
            p.add_run("Scenario: ").bold = True
            p.add_run(sc)
            doc.add_paragraph("Your script:").runs[0].bold = True
            add_ruled_lines(doc, 2)

    elif at in ("annotation_practice",):
        doc.add_paragraph("Write your 2-sentence summary here:").runs[0].bold = True
        add_ruled_lines(doc, 2)

    elif at == "timed_write":
        p = doc.add_paragraph()
        p.add_run(f"Time: {act['minutes']} min  ·  Plan (2 min): __________________  ·  Draft:")
        p.runs[0].bold = True
        add_ruled_lines(doc, act.get("lines", 16))

    elif at == "checklist":
        for item in act["items"]:
            doc.add_paragraph(f"☐  {item}").runs[0].font.size = Pt(10)

    elif at == "gpa_calculator":
        t = doc.add_table(rows=1 + len(act["courses"]) + 1, cols=5)
        t.style = "Table Grid"
        for j, hdr in enumerate(["Course", "Grade", "Credits", "GPA Points", "Weighted Points"]):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(hdr).font.color.rgb = WHITE
        for i, c in enumerate(act["courses"], 1):
            t.rows[i].cells[0].paragraphs[0].add_run(c[0]).font.size = Pt(9)
            t.rows[i].cells[1].paragraphs[0].add_run(c[1]).font.size = Pt(9)
            t.rows[i].cells[2].paragraphs[0].add_run(c[2]).font.size = Pt(9)
        set_cell_bg(t.rows[-1].cells[2], WARRIOR_GOLD)
        t.rows[-1].cells[2].paragraphs[0].add_run("TOTAL →").font.size = Pt(9)

    elif at == "habits_checklist":
        for item in act["items"]:
            doc.add_paragraph(f"☐  {item}").runs[0].font.size = Pt(10)

    elif at == "commitment_card":
        for f in act["fields"]:
            doc.add_paragraph(f, style="List Bullet")
            add_ruled_lines(doc, 1)

    elif at == "activities_checklist":
        for cat, items in act["categories"].items():
            doc.add_paragraph(cat).runs[0].bold = True
            t = doc.add_table(rows=(len(items) + 2) // 3, cols=3)
            t.style = "Table Grid"
            row_i, col_i = 0, 0
            for item in items:
                t.rows[row_i].cells[col_i].paragraphs[0].add_run(f"☐  {item}").font.size = Pt(9)
                col_i += 1
                if col_i == 3:
                    col_i = 0
                    row_i += 1

    elif at == "time_map":
        cols = act["columns"]
        t = doc.add_table(rows=1 + len(act["hours"]), cols=1 + len(cols))
        t.style = "Table Grid"
        set_cell_bg(t.rows[0].cells[0], WARRIOR_DARK)
        t.rows[0].cells[0].paragraphs[0].add_run("Time").font.color.rgb = WHITE
        for j, c in enumerate(cols, 1):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(c).font.color.rgb = WHITE
        for i, h in enumerate(act["hours"], 1):
            label = f"{h % 12 or 12}:00 {'AM' if h < 12 else 'PM'}"
            t.rows[i].cells[0].paragraphs[0].add_run(label).font.size = Pt(8)

    elif at in ("schedule_grid", "full_schedule_builder"):
        raw_sems = act.get("semesters", ["Semester 1", "Semester 2"])
        sems = raw_sems if isinstance(raw_sems, list) else ["Semester 1", "Semester 2"]
        periods = act.get("periods", act.get("periods", 7))
        if isinstance(periods, list):
            period_labels = periods
        else:
            period_labels = [f"Period {p}" for p in range(1, periods + 1)]
        t = doc.add_table(rows=1 + len(period_labels), cols=3)
        t.style = "Table Grid"
        set_cell_bg(t.rows[0].cells[0], WARRIOR_DARK)
        t.rows[0].cells[0].paragraphs[0].add_run("Period").font.color.rgb = WHITE
        for j, sem in enumerate(sems[:2], 1):
            set_cell_bg(t.rows[0].cells[j], WARRIOR_DARK)
            t.rows[0].cells[j].paragraphs[0].add_run(sem).font.color.rgb = WHITE
        for i, pl in enumerate(period_labels, 1):
            t.rows[i].cells[0].paragraphs[0].add_run(pl).font.size = Pt(9)
        if act.get("note"):
            doc.add_paragraph(act["note"]).runs[0].italic = True
        if act.get("include_reflection"):
            doc.add_paragraph("Reflection: Which elective connects to your Unit 2 career pathway? Explain:").runs[0].bold = True
            add_ruled_lines(doc, 3)

    elif at == "year_plan_summary":
        for label, lines in act["fields"]:
            doc.add_paragraph(label).runs[0].bold = True
            add_ruled_lines(doc, lines)

    elif at == "final_schedule_with_rationale":
        t = doc.add_table(rows=1 + act["periods"], cols=4)
        t.style = "Table Grid"
        for j, hdr in enumerate(["Period", "Course Choice", "Why I chose it", "Alternate"]):
            set_cell_bg(t.rows[0].cells[j], uc)
            t.rows[0].cells[j].paragraphs[0].add_run(hdr).font.color.rgb = WHITE
        for i in range(1, act["periods"] + 1):
            t.rows[i].cells[0].paragraphs[0].add_run(f"Period {i}").font.size = Pt(9)

    elif at in ("question_prep",):
        for i in range(1, act["num_questions"] + 1):
            doc.add_paragraph(f"Question {i}:").runs[0].bold = True
            add_ruled_lines(doc, act["lines_per_question"])

    elif at == "panel_notes":
        add_ruled_lines(doc, act["lines"])


def add_exit_ticket(doc, unit, lesson):
    uc = color_from_tuple(unit["color"])
    et = lesson["exit_ticket"]
    et_type = et.get("type", "quick_write")

    et_t = doc.add_table(rows=2, cols=1)
    et_t.style = "Table Grid"
    set_cell_bg(et_t.rows[0].cells[0], uc)
    r = et_t.rows[0].cells[0].paragraphs[0].add_run(
        f"✏  EXIT TICKET  ·  Lesson {lesson['lesson_label']}: {lesson['title']}")
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

    bc = et_t.rows[1].cells[0]
    set_cell_bg(bc, CREAM_BG)

    if et_type in ("quick_write", "synthesis_write"):
        bc.add_paragraph(et["prompt"]).runs[0].font.size = Pt(10)
        for _ in range(et.get("lines", 5)):
            bc.add_paragraph("")

    elif et_type == "three_two_one":
        for p in et["prompts"]:
            pp = bc.add_paragraph(p)
            pp.runs[0].italic = True
            pp.runs[0].font.size = Pt(10)
            for _ in range(2):
                bc.add_paragraph("")

    elif et_type == "reflection_scale":
        bc.add_paragraph(et["prompt"]).runs[0].font.size = Pt(10)
        sc_t = doc.add_table(rows=1, cols=5)
        sc_t.style = "Table Grid"
        for i, label in enumerate(["1 — Not at all", "2", "3 — Somewhat", "4", "5 — Absolutely"]):
            sc_t.rows[0].cells[i].paragraphs[0].add_run(label).font.size = Pt(9)
        bc.add_paragraph("")
        for _ in range(et.get("lines", 3)):
            bc.add_paragraph("")

    elif et_type == "sentence_stems":
        for p in et["prompts"]:
            bc.add_paragraph(p).runs[0].italic = True
            for _ in range(2):
                bc.add_paragraph("")

    elif et_type == "vocabulary_check":
        bc.add_paragraph(et["prompt"]).runs[0].font.size = Pt(10)
        for term, defn in zip(et["terms"], et["definitions"]):
            pp = bc.add_paragraph()
            pp.add_run(f"{term}  ↔  ").bold = True
            pp.add_run(defn).font.size = Pt(9)
        bc.add_paragraph("Write TWO terms in sentences:").runs[0].italic = True
        for _ in range(3):
            bc.add_paragraph("")

    elif et_type == "planning_response":
        for p in et["prompts"]:
            bc.add_paragraph(p).runs[0].italic = True
            for _ in range(2):
                bc.add_paragraph("")

    elif et_type == "postcard":
        bc.add_paragraph(et["prompt"]).runs[0].font.size = Pt(10)
        for _ in range(et.get("lines", 5)):
            bc.add_paragraph("")

    else:
        bc.add_paragraph(str(et.get("prompt", ""))).runs[0].font.size = Pt(10)
        for _ in range(4):
            bc.add_paragraph("")

    doc.add_page_break()


def build():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    add_cover(doc)
    add_how_to_use(doc)
    add_planner(doc)

    for unit in UNITS:
        add_unit_title(doc, unit)
        add_test_page(doc, unit, is_pretest=True)
        for lesson in unit["lessons"]:
            add_cornell_page(doc, unit, lesson)
            add_lesson_content(doc, unit, lesson)
            add_exit_ticket(doc, unit, lesson)
        add_test_page(doc, unit, is_pretest=False)

    out = "output/warrior_prep_student_workbook.docx"
    doc.save(out)
    print(f"Word doc saved → {out}")


if __name__ == "__main__":
    build()
